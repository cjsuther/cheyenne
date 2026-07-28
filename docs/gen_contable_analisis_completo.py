#!/usr/bin/env python3
"""Genera el Word con el análisis funcional COMPLETO del sistema legacy CONTABLE (11 módulos)."""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(__file__)
PRIMARY = RGBColor(0x1D, 0x4E, 0xD8)
DARK = RGBColor(0x1F, 0x29, 0x37)
GREY = RGBColor(0x6B, 0x72, 0x80)

doc = Document()
n = doc.styles["Normal"]; n.font.name = "Calibri"; n.font.size = Pt(11)
n.paragraph_format.space_after = Pt(6)


def shade(cell, hexc):
    tcPr = cell._tc.get_or_add_tcPr(); sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear"); sh.set(qn("w:color"), "auto"); sh.set(qn("w:fill"), hexc)
    tcPr.append(sh)


def h1(t):
    p = doc.add_heading(t, level=1)
    for r in p.runs: r.font.color.rgb = PRIMARY
    return p


def h2(t):
    p = doc.add_heading(t, level=2)
    for r in p.runs: r.font.color.rgb = DARK
    return p


def para(t="", bold=False, italic=False, color=None):
    p = doc.add_paragraph(); r = p.add_run(t)
    r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = color
    return p


def bullet(t, lvl=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + 0.2 * lvl)
    p.add_run(t); return p


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
    for i, htext in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""
        run = c.paragraphs[0].add_run(htext); run.bold = True; run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); shade(c, "1D4ED8")
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""; run = cells[i].paragraphs[0].add_run(str(v)); run.font.size = Pt(9)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    return t


def note(t):
    p = doc.add_paragraph(); r = p.add_run("  ⚠  " + t); r.italic = True; r.font.color.rgb = RGBColor(0x92, 0x40, 0x0E)
    pPr = p._p.get_or_add_pPr(); sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear"); sh.set(qn("w:color"), "auto"); sh.set(qn("w:fill"), "FEF3C7"); pPr.append(sh)
    return p


# ── Portada ──────────────────────────────────────────────────────────
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Sistema CONTABLE"); r.bold = True; r.font.size = Pt(30); r.font.color.rgb = PRIMARY
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run("Análisis funcional completo del código fuente"); r.font.size = Pt(16); r.font.color.rgb = DARK
s2 = doc.add_paragraph(); s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s2.add_run("ERP de Administración Financiera Municipal (RAFAM) · VB6 + SQL Server · Ver 12.x · 11 módulos · 589 reportes")
r.italic = True; r.font.color.rgb = GREY
doc.add_paragraph()

# ── 1. Resumen ejecutivo ─────────────────────────────────────────────
h1("1. Resumen ejecutivo")
para("El sistema \"CONTABLE\" es un ERP completo de Administración Financiera y de Recursos de los "
     "Municipios (RAFAM): administra el presupuesto, el ciclo del gasto, las compras y el stock, la "
     "tesorería (recaudación y pagos), la deuda pública, el patrimonio, la inversión pública y la "
     "rendición de cuentas regulatoria. Está construido en Visual Basic 6 sobre SQL Server, con "
     "reportería Crystal Reports (589 reportes) y cubos OLAP para el tablero gerencial.")
bullet("11 módulos ejecutables + 1 librería común (vbComun, 626 componentes) + 12 DLLs de framework del proveedor (CB1/CB5).")
bullet("Mismo proveedor y framework que el sistema legacy de Rentas: comparten seguridad, personas, tributos, cuentas contables y base de datos (integración nativa Contable ↔ Rentas).")
bullet("Cumplimiento regulatorio RAFAM (provincia de Buenos Aires): 28 reportes en formato oficial y generación de la rendición de cuentas trimestral/anual.")
bullet("2.672 archivos · 820 MB de fuente · versión 12.7 (activo y mantenido).")

# ── 2. Arquitectura técnica ──────────────────────────────────────────
h1("2. Arquitectura técnica")
h2("2.1 Composición")
table(["Pieza", "Contenido"], [
    ["11 módulos VB6", "Cada uno con dos proyectos: X.vbp (ejecutable de UI, MDI con menús) y XEC.vbp (capa de entidades de negocio)"],
    ["vbComun (38 MB)", "626 componentes compartidos: comprobantes (con adjuntos y firma digital), asientos, movimientos, proveedores, extracontables, presupuesto, stock, login, ejercicios, bloqueos, auditoría"],
    ["Framework CB (12 DLLs)", "cb1dbADO (acceso a datos/SQL), cb5Segur (seguridad), cb5Pers (personas), cb5CtasCont (cuentas contables), cb5trita (tributos y tasas), cb5AppRe (Application Rentas), cb5Obras, cb5Saldo, cb5reten, cb5caloc, configuracion, calculoDeuda"],
    ["CReports (723 MB)", "589 reportes Crystal Reports organizados por módulo"],
], widths=[1.7, 5.2])
h2("2.2 Datos y transversales")
bullet("SQL Server (sqloledb). Config relevada: server CUBIX, BD principal BdMatanza y BD espejo de auditoría AudMatanza (producto multi-municipio: la config muestra clientes Tigre/La Matanza).")
bullet("Doble conexión por proceso (escrituras vs. lecturas) para evitar bloqueos; límite de filas y timeouts parametrizados en Contable.ini.")
bullet("Auditoría dual: cada cambio se compara campo a campo contra la BD espejo y se registra con acción (I/U/D), usuario y fecha.")
bullet("Multi-ejercicio: todas las tablas llevan el año fiscal; el usuario elige el ejercicio al ingresar; aperturas, cierres y prórrogas por año.")
bullet("Lógica de negocio en clases VB6 con SQL dinámico (los stored procedures son excepcionales) — favorable para una migración leyendo las clases.")
bullet("Seguridad por usuario → rol → recurso de subsistema (administrada por el módulo Administrador con cb5Segur).")
note("Hallazgo: credenciales de conexión a la BD en texto plano en Contable.ini (usuario CBLOGIN). "
     "La seguridad es a nivel aplicación, no de base de datos.")

# ── 3. Módulos ───────────────────────────────────────────────────────
h1("3. Funcionalidad por módulo")

h2("3.1 Contaduría — núcleo contable (130 formularios)")
para("Ciclo del gasto RAFAM completo con registro contable automático. Detallado en el documento "
     "\"Módulo Contaduría — Funcionalidad\"; síntesis:", italic=True, color=GREY)
bullet("Ciclo: Solicitud → Preventivo → Compromiso (con/sin OC) → Devengado → Orden de Pago → Liquidación (con retenciones) → Pagado; más cancelación de deuda, reintegros, regularizaciones y sueldos.")
bullet("Contabilidad: plan de cuentas, asientos por minuta y automáticos, cierre mensual, apertura/cierre de ejercicio (deuda flotante, resultados, cierre patrimonial), libros Diario/Mayor, balances y estados patrimoniales/financieros.")
bullet("Recursos: devengamiento y devoluciones; ejecución por rubro/carácter/procedencia.")
bullet("Retenciones: IIBB 84/08, Ganancias, IVA, SIJP, AFIP-SICOSS con archivos TXT regulatorios.")
bullet("Extracontables (fondos de terceros), embargos judiciales (carga TXT), cesiones de crédito, garantías.")
bullet("Conciliación bancaria; rendición de cuentas RAFAM en PDF (trimestral/semestral/anual) con anulación y re-emisión; hojas rubricadas.")
bullet("Consultas de gestión: ejecución presupuestaria y extrapresupuestaria, deuda flotante, cuentas corrientes de proveedores, memoria anual.")

h2("3.2 Tesorería — caja y bancos (92 clases)")
bullet("Recaudación: cajas registradoras con cierre centralizado, tickets, recaudaciones manuales y automáticas por lotes (débito de sueldos, Interbanking, EPagos, MercadoPago, IntGO Pagos, rendición genérica), importación de archivos, búsqueda de recibos, cierre y exportación con control de balanceados.")
bullet("Egresos: órdenes de pago, cheques sobre órdenes (con numeración y correlatividad), cheques anulados/no presentados/pago diferido con proceso al vencimiento, órdenes bancarias con generación de archivo, registro del pagado, órdenes de reintegro.")
bullet("Retenciones impositivas: débitos y depósitos, individuales y grupales, con reemplazo de débito.")
bullet("Transferencias bancarias (simultáneas, origen, depósito); depósitos; otros ingresos; comisiones; anulación y regularización de ingresos; conciliación de saldo de caja (manual/automática).")
bullet("Conciliación bancaria (extractos, movimientos anteriores, apertura de ejercicio).")
bullet("Registros: beneficiarios y poderes con cuentas bancarias, cesiones de crédito, embargos (estados y juzgados), escribanos, proveedores con tipos de autorización, cuentas y entidades bancarias, bocas recaudadoras, cajas, clearing, monedas, fondos afectados y de terceros.")
bullet("Programación de caja (F47) y deuda exigible del ejercicio (F48); garantías; bandeja de firma.")
bullet("Libros: parte diario de ingresos y de egresos, caja, banco, composición del saldo de caja, resumen de saldos de tesorería; más de 15 consultas multifiltro (cheques, depósitos, transferencias, órdenes).")

h2("3.3 Presupuesto — ciclo presupuestario (54 clases)")
bullet("Formulación (formularios F1 a F11): política presupuestaria, programación de recursos, estructura programática y descripción de programas, metas por programa, recursos humanos por objeto y por cargo, compras de bienes y servicios, programación física y financiera de proyectos, gestión de la deuda, transferencias.")
bullet("Modificaciones: al presupuesto de gastos, al cálculo de recursos, y movimientos de personal/cargos (F72).")
bullet("Ejecución: programación de metas y recursos, programación financiera del compromiso, cuotas de compromiso por jurisdicción/fuente/inciso, programación física y financiera de proyectos (F14/F15).")
bullet("Evaluación: causas de desvíos operativas, financieras y técnicas; conclusiones y recomendaciones.")
bullet("Consultas multidimensión: por inciso, objeto, fuente, categoría programática, finalidad y función, carácter económico, jurisdicción, con perspectiva de género; cuenta Ahorro-Inversión-Financiamiento; exportación de datos.")
bullet("Nomencladores: jurisdicciones, rubros, objetos del gasto, fuentes, carácter económico, finalidad y función, unidades ejecutoras, metas, unidades de medida, proyectos (clases, etapas, prioridades), bloqueo presupuestario, cargos y categorías.")

h2("3.4 Crédito Público — deuda pública (45 clases)")
bullet("Empréstitos: alta y categorización (bilateral/multilateral), participantes y sus relaciones, montos históricos.")
bullet("Cronogramas: amortización del capital, intereses en período de gracia, intereses de financiamiento, intereses por mora, comisiones y gastos.")
bullet("Desembolsos y colocaciones estimados vs. reales; instrucciones de pago; pagos reales.")
bullet("Multimoneda con tipos de cambio; tipos de interés (simple/compuesto); medios de desembolso y de pago.")

h2("3.5 Compras — adquisiciones")
bullet("Proveedores: registro con ficha, pre-inscripción, tipos, estados, sociedades, referencias, consejos profesionales, tributos.")
bullet("Artículos: catálogo con agrupamientos, unidades y precios de referencia.")
bullet("Proceso de compra: pedidos de cotización → cotizaciones → acta de apertura → comparativo → adjudicación → orden de compra → recepción → factura/comprobante → remitos; devolución de materiales.")
bullet("Consultas: OC por oficina/proveedor/código postal/anuladas, preventivas sin OC, licitaciones, facturas conformes, garantías, ficha de stock y stock por OC.")

h2("3.6 Depósito — almacenes y stock")
bullet("Recepciones de órdenes de compra; comprobantes y remitos.")
bullet("Stock por depósito y artículo con ficha de stock; consumos generales y de combustible (YPF); consumos extracontables.")
bullet("Transferencias entre depósitos; ajustes de stock; devoluciones a depósito y a proveedor.")
bullet("Generación masiva de PDFs; bandeja de firma; consultas de pendientes por OC y avances de comprobantes.")

h2("3.7 Solicitante — pedidos de las áreas")
bullet("Pre-solicitudes (borrador) y solicitudes de pedido de cada área, con depuración y consultas multifiltro.")
bullet("Bandeja de firma para aprobación jerárquica.")
bullet("Seguimiento de recepciones y estado de entregas (el área ve en qué está su pedido).")

h2("3.8 Patrimonio — bienes de uso (36 clases)")
bullet("Bienes de dominio privado: inventario, altas por orden de compra o recepción, bajas por cuenta, cambio de oficina, pases entre responsables con remitos.")
bullet("Bienes de dominio público: registro con nomenclatura catastral, responsables, relación con la contabilidad (CGP), altas y bajas.")
bullet("Amortización: sistemas de valuación y de amortización parametrizables, cálculo periódico.")
bullet("Maestros: estados del bien, motivos de alta/baja, responsables, cuentas patrimoniales, oficinas.")

h2("3.9 Inversión Pública — banco de proyectos (166 clases, 77 reportes)")
bullet("Banco de proyectos con ciclo completo: captación (F.58 A-E), formulación (F.59-F.65), selección (F.66-F.68), ejecución (F.69-F.70) y evaluación (F.71).")
bullet("Formularios de identificación y evaluación I.1 a I.12 (erogaciones, RRHH, cronogramas, evaluación económica y financiera, costo-beneficio, evaluación ex-post).")
bullet("Informes de cierre: ITP (terminación de proyecto, cuadros 1-3) e IOP (operación, cuadros 4-7); informes de desempeño trimestrales/anuales y rentabilidad.")
bullet("Seguimiento físico-financiero de actividades; impactos, medidas mitigatorias, incumplimientos de plazo, modificaciones de obra.")

h2("3.10 Gerencial — tablero de mando")
bullet("Consultas OLAP sobre cubos de Analysis Services: saldos contables/presupuestarios, inmuebles, comercios, operativos.")
bullet("Cruza el mundo contable con el de Rentas (inmuebles, comercios, cementerios, planes de pago) y con sueldos, expedientes, vehículos y plazos fijos.")
bullet("Exportación a Excel con tablas dinámicas; filtros por período/cuota/localidad.")

h2("3.11 Administrador — seguridad y parámetros")
bullet("Usuarios, roles y accesos; recursos protegibles organizados en árbol y asignados por subsistema; restricciones por usuario; funcionarios responsables.")
bullet("Parámetros del sistema, numeraciones de comprobantes, operaciones permitidas, acrónimos, feriados, ejercicios.")
bullet("Informes de auditoría de accesos (usuarios/accesos, recursos/subsistema).")

# ── 4. Circuito integrado ────────────────────────────────────────────
h1("4. El circuito integrado del gasto")
para("Los módulos no son islas: implementan un único flujo de negocio con control presupuestario y "
     "registro contable en cada paso.")
table(["Paso", "Módulo", "Resultado"], [
    ["1. El área pide", "Solicitante", "Solicitud de pedido aprobada (bandeja de firma)"],
    ["2. Se reserva crédito", "Contaduría", "Preventivo sobre la partida presupuestaria"],
    ["3. Se cotiza y adjudica", "Compras", "Cotizaciones, comparativa, adjudicación (licitación si corresponde)"],
    ["4. Se compromete", "Contaduría", "Compromiso con la orden de compra"],
    ["5. Se recibe", "Depósito", "Recepción, entrada a stock (o alta patrimonial en Patrimonio)"],
    ["6. Se devenga", "Contaduría", "Obligación de pago con la factura conformada"],
    ["7. Se ordena y liquida", "Contaduría", "Orden de pago + retenciones impositivas"],
    ["8. Se paga", "Tesorería", "Cheque / transferencia / orden bancaria; registro del pagado"],
    ["9. Se contabiliza y rinde", "Contaduría", "Asientos automáticos, cierres, rendición RAFAM"],
], widths=[1.7, 1.3, 3.9])
para("En paralelo: Presupuesto define el marco (crédito y cuotas de compromiso); Crédito Público e "
     "Inversión Pública administran sus casos especiales (deuda y proyectos); Gerencial observa todo; "
     "Administrador da seguridad transversal.", italic=True, color=GREY)

# ── 5. Reportería ────────────────────────────────────────────────────
h1("5. Reportería (589 Crystal Reports)")
table(["Carpeta", "Cant.", "Contenido representativo"], [
    ["Contaduria", "192", "Balance General, Sumas y Saldos, Libro Diario/Mayor, cheques, órdenes de pago (variantes RAFAM/ARBA), retenciones, ejecución presupuestaria"],
    ["Compras", "96", "Licitaciones y actas, comparativas, órdenes de compra (RAFAM), recepciones, fichas de proveedor"],
    ["InversionPublica", "77", "Formularios I1-I12/FI10, indicadores ITP/IOP, hojas de ruta"],
    ["Presupuestos", "76", "Presupuesto por objeto/fuente/programa, ejecución física y financiera, matriz de recursos"],
    ["Tesoreria", "68", "Partes diarios, balance de tesorería, recaudación por origen/lugar, embargos, transferencias"],
    ["Patrimonio", "30", "Altas/bajas, constancias, resumen de cuenta, amortización"],
    ["CreditoPublico", "21", "Empréstitos, intereses, amortización, desembolsos"],
    ["Depositos", "18", "Stock, consumos, transferencias, devoluciones"],
    ["Administrador", "9", "Accesos, recursos, usuarios (RAFAM)"],
], widths=[1.5, 0.6, 4.8])
bullet("28 reportes en formato RAFAM (obligatorio regulatorio) y variantes ARBA.")
bullet("Variantes por formato (A4, horizontal/vertical, detallado/resumido/sintético); pocos obsoletos (old/new).")

# ── 6. Integraciones externas ────────────────────────────────────────
h1("6. Integraciones externas")
bullet("Sistema Rentas (Ingresos Públicos): mismo framework, BD compartida y conexión directa; los ingresos tributarios ejecutan el cálculo de recursos y el módulo Gerencial cruza ambos mundos.")
bullet("AFIP: web services de comprobantes, SICOSS y registro de compras (TXT).")
bullet("ARBA: retenciones IIBB (Normativa 84/08) y formatos de órdenes de pago.")
bullet("Pasarelas de pago: EPagos, MercadoPago, IntGO, Interbanking, débito de sueldos.")
bullet("Bancos: archivos de órdenes bancarias, extractos para conciliación, clearing.")
bullet("Firma digital de documentos (SHA-256, bandejas de firma en varios módulos).")

# ── 7. Relación con Cheyenne ─────────────────────────────────────────
h1("7. Relación con Cheyenne y camino de adopción")
para("El super-módulo \"Contaduría\" creado en Cheyenne es el lugar natural para absorber este "
     "dominio. Puntos de partida:")
bullet("Cheyenne ya tiene la mitad \"ingresos\" de Tesorería (recaudación por lotes) y los maestros contables en Administración/Configuración (cuentas contables, recursos por rubro, jurisdicciones).")
bullet("La integración Contable ↔ Rentas del legacy equivale en Cheyenne a contaduría ↔ ingresos_publicos (ya comunicados por HTTP).")
bullet("La lógica reside en clases VB6 legibles (sin stored procedures): misma técnica de migración usada con el motor de cálculo de Rentas.")
h2("Fases sugeridas")
table(["Fase", "Alcance", "Fuente legacy"], [
    ["1", "Núcleo contable: plan de cuentas, ejercicios, asientos, libros + ciclo del gasto (preventivo → pagado)", "vbContaduria"],
    ["2", "Circuito de adquisiciones: solicitudes, cotización/adjudicación, OC, recepciones, stock", "vbSolicitante + vbCompras + vbDeposito"],
    ["3", "Egresos de tesorería: cheques, órdenes bancarias, transferencias, conciliación", "vbTesoreria (egresos)"],
    ["4", "Presupuesto: formulación, modificaciones, ejecución con cuotas", "vbPresupuesto"],
    ["5", "Patrimonio, Crédito Público, Inversión Pública y tablero gerencial", "restantes"],
], widths=[0.6, 3.6, 2.7])
note("Restricción regulatoria: los formatos RAFAM (reportes y rendición de cuentas) son requisito "
     "legal — cualquier reemplazo debe reproducirlos con exactitud. La reportería (589 reportes) es "
     "el mayor esfuerzo individual de la migración.")

para()
para("Documento generado a partir del relevamiento del código fuente en legacy/CONTABLE/Ver12.0 "
     "(menús MDI, clases de negocio, tablas SQL, proyectos VBP, configuración y reportes). "
     "Complementa: \"Módulo Contaduría — Funcionalidad\" y legacy/analisis/ANALISIS-CONTABLE.md.",
     italic=True, color=GREY)

out = os.path.join(HERE, "Contable-Analisis-Funcional-Completo.docx")
doc.save(out)
print("OK ->", out)
