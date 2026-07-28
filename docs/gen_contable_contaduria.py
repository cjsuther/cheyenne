#!/usr/bin/env python3
"""Genera el Word funcional del MÓDULO CONTADURÍA (solo contabilidad) del legacy CONTABLE."""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(__file__)
PRIMARY = RGBColor(0x1D, 0x4E, 0xD8)
DARK = RGBColor(0x1F, 0x29, 0x37)
GREY = RGBColor(0x6B, 0x72, 0x80)

doc = Document()
n = doc.styles["Normal"]; n.font.name = "Calibri"; n.font.size = Pt(11)
n.paragraph_format.space_after = Pt(6)


def shade(cell, hexc):
    tcPr = cell._tc.get_or_add_tcPr(); sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear"); sh.set(qn("w:color"), "auto"); sh.set(qn("w:fill"), hexc)
    tcPr.append(sh)


def h1(t):
    p = doc.add_heading(t, level=1)
    for r in p.runs: r.font.color.rgb = PRIMARY
    return p


def h2(t):
    p = doc.add_heading(t, level=2)
    for r in p.runs: r.font.color.rgb = DARK
    return p


def para(t="", bold=False, italic=False, color=None):
    p = doc.add_paragraph(); r = p.add_run(t)
    r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = color
    return p


def bullet(t, lvl=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + 0.2 * lvl)
    p.add_run(t); return p


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
    for i, htext in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""
        run = c.paragraphs[0].add_run(htext); run.bold = True; run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); shade(c, "1D4ED8")
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""; run = cells[i].paragraphs[0].add_run(str(v)); run.font.size = Pt(9)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    return t


# ── Portada ──────────────────────────────────────────────────────────
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Sistema CONTABLE"); r.bold = True; r.font.size = Pt(30); r.font.color.rgb = PRIMARY
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run("Módulo Contaduría — Funcionalidad"); r.font.size = Pt(16); r.font.color.rgb = DARK
s2 = doc.add_paragraph(); s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s2.add_run("Qué tiene y qué realiza, a nivel funcional · relevado del código fuente VB6 (Ver 12.x)")
r.italic = True; r.font.color.rgb = GREY
doc.add_paragraph()

# ── 1. Propósito ─────────────────────────────────────────────────────
h1("1. Propósito y alcance")
para("El módulo Contaduría es el núcleo contable del ERP de administración financiera municipal "
     "(RAFAM). Registra y controla la ejecución del gasto y de los recursos del municipio, lleva la "
     "contabilidad general (asientos, libros, balances), administra los fondos de terceros "
     "(extracontables), practica las retenciones impositivas, concilia los bancos y produce la "
     "rendición de cuentas al Honorable Tribunal de Cuentas en los formatos RAFAM.")
para("Es el módulo más grande del sistema: 130 formularios, ~700 opciones de menú y más de 40 tablas "
     "propias, con integración directa a Tesorería, Compras, Presupuesto y Patrimonio.", italic=True, color=GREY)

# ── 2. Ciclo del gasto ───────────────────────────────────────────────
h1("2. El ciclo del gasto (columna vertebral del módulo)")
para("Toda erogación municipal recorre etapas encadenadas; cada etapa afecta el presupuesto y "
     "genera su registro contable:")
table(["Etapa", "Qué significa", "Funciones que ofrece el módulo"], [
    ["1. Solicitud de Pedido", "Un área pide un gasto", "Alta de solicitud, depuración, consultas multifiltro"],
    ["2. Preventivo", "Reserva del crédito presupuestario", "Solicitud de gastos; ajustes y desafectaciones; reimputación de preventivo de ejercicios anteriores; desafectación general"],
    ["3. Compromiso", "Afectación firme de la partida", "Con o sin orden de compra; reimputación de OC; desafectaciones; proveedores con compromisos pendientes; control OC vs. recepción; pendientes sin devengar"],
    ["4. Devengado", "Nace la obligación de pago (llegó la factura/recepción)", "Devengado y desafectaciones; recepción de comprobantes; facturas devengadas"],
    ["5. Orden de Pago", "Se ordena pagar", "OP presupuestarias, extrapresupuestarias y de comisiones; impresión por rangos; etiquetas; registro"],
    ["6. Liquidación de Pago", "Se liquida al beneficiario, con retenciones", "Liquidaciones de pago y de comisiones; reliquidación de ejercicios anteriores; cálculo de retenciones"],
    ["7. Pagado", "Tesorería ejecuta el pago", "Registro del Pagado (cierre del circuito, cheques/transferencias en Tesorería)"],
], widths=[1.4, 1.9, 3.6])
bullet("Cancelación de deuda: órdenes de cancelación de deuda de ejercicios anteriores, con reliquidación.")
bullet("Órdenes de reintegro: reingreso de fondos sobre órdenes de pago ya emitidas.")
bullet("Regularizaciones: correcciones de gastos/recursos presupuestarios y extrapresupuestarios, y cambio de imputación.")
bullet("Sueldos: generación automática e importación de la liquidación de haberes al circuito del gasto.")

# ── 3. Contabilidad general ──────────────────────────────────────────
h1("3. Contabilidad general")
h2("3.1 Plan de cuentas y maestros contables")
bullet("Plan de Cuentas con tipos de cuenta y agrupamientos (para reportes jerárquicos).")
bullet("Cuentas patrimoniales; fondos afectados; fondos de terceros y especiales.")
bullet("Clasificadores presupuestarios: objetos del gasto, recursos por rubro, estructura programática, finalidad y función, carácter económico, fuente de financiamiento.")
h2("3.2 Asientos")
bullet("Asientos por minuta (manuales) y asientos automáticos generados por las transacciones del ciclo del gasto.")
bullet("Tipos de asiento parametrizables; numeración y correlatividad de comprobantes.")
bullet("Cierre mensual (bloqueo del período).")
h2("3.3 Apertura y cierre de ejercicio")
bullet("Apertura de ejercicio: asiento patrimonial, cuenta corriente, prórroga del cálculo de recursos y del presupuesto de gastos.")
bullet("Cierre anual: pagos y altas de deuda flotante, cierre de recursos y de gastos del ejercicio, resultado del ejercicio y de ejercicios anteriores, cierre patrimonial.")
bullet("Parámetros por ejercicio y trabajo multi-ejercicio (cada registro lleva el año fiscal).")
h2("3.4 Libros y estados")
bullet("Libro Diario y Libro Mayor.")
bullet("Balance de Sumas y Saldos y Balance General.")
bullet("Estado de Situación Patrimonial; Evolución del Activo, del Pasivo y combinada.")
bullet("Disponibilidades; Cuenta Ahorro-Inversión-Financiamiento; Situación Económico-Financiera.")
bullet("Estado de afectación de saldos; Resumen presupuestario, financiero y patrimonial; Resultado Art. 44.")
bullet("Cuentas corrientes de proveedores (detallada, resumen y extrapresupuestaria).")

# ── 4. Recursos ──────────────────────────────────────────────────────
h1("4. Recursos (ingresos)")
bullet("Devengamiento de recursos (reconocimiento del derecho a cobrar).")
bullet("Devoluciones de recursos, con impresión por rangos, etiquetas y registro.")
bullet("Ejecución de recursos por rubro, carácter económico y procedencia, con detalle.")

# ── 5. Retenciones ───────────────────────────────────────────────────
h1("5. Retenciones e información impositiva")
bullet("Tipos de retenciones parametrizables; clave CIT; retención individual y por rangos.")
bullet("IIBB Normativa 84/08 (ARBA): proceso, consultas e informe.")
bullet("Generación de archivos TXT regulatorios: Ingresos Brutos, Ganancias, IVA, SIJP, AFIP-SICOSS y Registro de Compras AFIP.")
bullet("Consultas de retenciones ingresadas y liquidadas; IVA compras.")

# ── 6. Extracontables ────────────────────────────────────────────────
h1("6. Extracontables (fondos de terceros)")
para("Movimientos que no ejecutan presupuesto pero pasan por las cuentas del municipio: fondos de "
     "terceros, depósitos en garantía, retenciones a depositar, etc.")
bullet("Registro extracontable por comprobante, por objeto y por patrimonio.")
bullet("Ejecución extrapresupuestaria: estados de recursos y gastos, detalle e ingresos/egresos.")
bullet("Órdenes de pago extrapresupuestarias y cuentas corrientes extrapresupuestarias.")

# ── 7. Embargos, cesiones y garantías ────────────────────────────────
h1("7. Embargos, cesiones de crédito y garantías")
bullet("Embargos sobre derechos de crédito de proveedores: carga por TXT judicial, consulta y liquidación del embargo en el pago.")
bullet("Cesiones de crédito: el proveedor cede el cobro a un tercero (cesionarios, estados).")
bullet("Garantías de proveedores: alta, tipos, valores y autorización a devolver.")

# ── 8. Conciliación bancaria ─────────────────────────────────────────
h1("8. Conciliación bancaria")
bullet("Carga de extractos bancarios y movimientos anteriores.")
bullet("Conciliación de cada cuenta bancaria; apertura de ejercicio de la conciliación; consultas.")

# ── 9. Rendición de cuentas (RAFAM) ─────────────────────────────────
h1("9. Rendición de cuentas (RAFAM)")
para("Producción regulatoria hacia el Honorable Tribunal de Cuentas de la provincia:")
bullet("Generación de la rendición en PDF por período: 1º trimestre, 1º semestre, 3º trimestre, 2º semestre y anual.")
bullet("Anulación y re-generación de una rendición; visualización de rendiciones emitidas.")
bullet("Reportes en formato RAFAM: balances, estados financieros, órdenes de pago, ejecución (28 formatos identificados en la reportería).")
bullet("Hojas rubricadas y correlatividad de comprobantes (exigencias formales del Tribunal).")

# ── 10. Consultas de gestión ─────────────────────────────────────────
h1("10. Consultas y reportes de gestión")
bullet("Ejecución presupuestaria de gastos por objeto, finalidad y función, fuente, categoría programática y carácter económico; detalle de ejecución y del pagado; multifiltro; exportación de datos.")
bullet("Fondos afectados: ejecución, estado de gastos, recursos vs. gastos.")
bullet("Deuda: de ejercicios anteriores, deuda flotante por imputación y por proveedor, total.")
bullet("Proveedores: facturas, garantías, saldos, cuentas corrientes, comprobantes por proveedor.")
bullet("Gestión de gastos: presupuesto, evolución, mensual Anexo II, por concepto, economía/exceso, variaciones presupuestarias, información de gestión (detallado, mensual, evolución trimestral).")
bullet("Gestión de recursos: cálculo, evolución, mensual, comparativo mensual.")
bullet("Movimiento de fondos, memoria general anual, estado de saldos, seguimiento de trámites, resumen de tareas pendientes.")

# ── 11. Maestros ─────────────────────────────────────────────────────
h1("11. Maestros que administra")
table(["Grupo", "Contenido"], [
    ["Cuentas", "Plan de cuentas, tipos, fondos afectados, fondos de terceros y especiales, cuentas patrimoniales"],
    ["Clasificadores", "Objetos del gasto, recursos por rubro, estructura programática, índices y tipos de índices"],
    ["Organización", "Oficinas y su ubicación; obras (con tipos y relación persona-obra)"],
    ["Terceros", "Proveedores (con tributos, referencias, sociedades), personas (con todas sus tablas auxiliares), barrios"],
    ["Documentación", "Ejercicios, tipos de asientos, tipos de comprobantes, errores de proceso, parámetros por ejercicio"],
], widths=[1.6, 5.3])

# ── 12. Transversales ────────────────────────────────────────────────
h1("12. Aspectos transversales")
bullet("Firma digital: bandeja de entrada de firma para autorizar documentos del circuito.")
bullet("Adjuntos: comprobantes con archivos (PDF) y firmas asociadas.")
bullet("Auditoría: toda modificación queda espejada en una base de auditoría (comparación antes/después por usuario y fecha).")
bullet("Seguridad: acceso por usuario/rol/recurso administrado desde el módulo Administrador; cambio de ejercicio y de clave desde el propio menú.")
bullet("Control de concurrencia: bloqueos y limpieza de transacciones en proceso.")

# ── 13. Interfaces ───────────────────────────────────────────────────
h1("13. Interfaces con el resto del ERP")
table(["Módulo", "Intercambio"], [
    ["Tesorería", "Recibe las órdenes de pago liquidadas y devuelve el Pagado (cheques/transferencias); comparte conciliación bancaria"],
    ["Compras / Depósito / Solicitante", "El compromiso nace de la orden de compra; el devengado, de la recepción/factura"],
    ["Presupuesto", "El preventivo/compromiso consumen el crédito y las cuotas definidas en Presupuesto"],
    ["Patrimonio", "Las altas de bienes se vinculan al comprobante de compra y al asiento"],
    ["Rentas (Ingresos Públicos)", "Sistema hermano del mismo proveedor: comparte personas, tributos y recaudación (los ingresos ejecutan el cálculo de recursos)"],
], widths=[1.8, 5.1])

para()
para("Documento generado a partir del relevamiento del código fuente (menús MDI, clases de negocio, "
     "tablas y reportes) en legacy/CONTABLE/Ver12.0. Ver también: Cheyenne — Análisis funcional "
     "completo del sistema CONTABLE.", italic=True, color=GREY)

out = os.path.join(HERE, "Contable-Modulo-Contaduria-Funcional.docx")
doc.save(out)
print("OK ->", out)
