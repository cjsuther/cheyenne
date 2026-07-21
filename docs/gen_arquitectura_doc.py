#!/usr/bin/env python3
"""Genera el documento Word de arquitectura y comunicación entre módulos."""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(__file__)
PRIMARY = RGBColor(0x1D, 0x4E, 0xD8)
DARK = RGBColor(0x1F, 0x29, 0x37)
GREY = RGBColor(0x6B, 0x72, 0x80)
RED = RGBColor(0xB4, 0x1C, 0x1C)

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Calibri"; normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)


def shade(cell, hexc):
    tcPr = cell._tc.get_or_add_tcPr(); sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear"); sh.set(qn("w:color"), "auto"); sh.set(qn("w:fill"), hexc)
    tcPr.append(sh)


def h1(t):
    p = doc.add_heading(t, level=1)
    for r in p.runs: r.font.color.rgb = PRIMARY
    return p


def h2(t):
    p = doc.add_heading(t, level=2)
    for r in p.runs: r.font.color.rgb = DARK
    return p


def para(t="", bold=False, italic=False, color=None):
    p = doc.add_paragraph(); r = p.add_run(t)
    r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = color
    return p


def bullet(t, lvl=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + 0.2 * lvl)
    p.add_run(t); return p


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
    for i, htext in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""
        run = c.paragraphs[0].add_run(htext); run.bold = True; run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); shade(c, "1D4ED8")
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""; run = cells[i].paragraphs[0].add_run(str(v)); run.font.size = Pt(9)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    return t


def note(t):
    p = doc.add_paragraph(); r = p.add_run("  ⚠  " + t); r.italic = True; r.font.color.rgb = RGBColor(0x92, 0x40, 0x0E)
    pPr = p._p.get_or_add_pPr(); sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear"); sh.set(qn("w:color"), "auto"); sh.set(qn("w:fill"), "FEF3C7"); pPr.append(sh)
    return p


# ── Portada ──────────────────────────────────────────────────────────
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Cheyenne"); r.bold = True; r.font.size = Pt(32); r.font.color.rgb = PRIMARY
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run("Arquitectura de la solución y comunicación entre módulos"); r.font.size = Pt(15); r.font.color.rgb = DARK
s2 = doc.add_paragraph(); s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s2.add_run("Entendimiento de los módulos desarrollados · interacción · duplicación de información")
r.italic = True; r.font.color.rgb = GREY
doc.add_paragraph()

# ── 1. Resumen ejecutivo ─────────────────────────────────────────────
h1("1. Resumen ejecutivo")
para("Cheyenne es un sistema de Ingresos Públicos construido como microservicios: 11 módulos "
     "backend (FastAPI), cada uno en su contenedor, un frontend React (SPA) y Nginx como reverse "
     "proxy. La base es PostgreSQL 16 compartida, con un esquema/prefijo de tablas por módulo.")
para("Reglas de arquitectura:", bold=True)
bullet("La comunicación entre módulos es SOLO por HTTP; ningún módulo lee la base de datos de otro.")
bullet("seguridad es la única autoridad de autenticación (JWT) y autorización (RBAC).")
bullet("La auditoría se implementa como middleware no bloqueante en todos los módulos.")
para("En la práctica, el acoplamiento entre módulos es muy bajo: hoy existe un único llamado de "
     "negocio módulo-a-módulo (emisiones → ingresos_publicos para el padrón), más la validación "
     "de token que todos hacen contra seguridad. Esto es sano, pero deja algunos módulos "
     "aislados o con solapamiento de datos, que se detallan en la sección 6.", italic=True, color=GREY)

# ── 2. Diagrama ──────────────────────────────────────────────────────
h1("2. Diagrama de arquitectura")
img = os.path.join(HERE, "arquitectura.png")
if os.path.exists(img):
    doc.add_picture(img, width=Inches(6.6))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
para("Figura 1 — Capas (frontend, proxy, módulos de negocio, datos) y los únicos flujos de "
     "comunicación reales entre módulos.", italic=True, color=GREY)

# ── 3. Módulos ───────────────────────────────────────────────────────
h1("3. Módulos desarrollados")
para("Madurez estimada por cantidad de endpoints y por si tienen pantalla en el frontend.")
table(["Módulo", "Rol", "Endpoints", "Frontend", "Madurez"], [
    ["ingresos_publicos", "Núcleo de Rentas: contribuyentes, cuentas, inmuebles/comercios/vehículos, base imponible, tasas, planes, padrón", "125", "Sí", "Alta"],
    ["administracion", "Maestros (personas, localidades, jurisdicciones, cuentas contables) + mesa de entradas / expedientes", "94", "Sí", "Alta"],
    ["tesoreria", "Recaudación, caja, recibos y rendiciones", "71", "Sí", "Alta"],
    ["emisiones", "Motor de cálculo, workflow de 16 pasos, liquidaciones, cuenta corriente, comprobantes, fórmulas", "35", "Sí", "Media-Alta"],
    ["seguridad", "Autenticación (JWT) y autorización (RBAC): usuarios, perfiles, permisos", "32", "Sí", "Alta"],
    ["wav", "Tributo por DD.JJ. con cuentas/DD.JJ./planes propios", "12", "Oculto", "Parcial"],
    ["interface", "Boletas / consulta web / notificaciones al contribuyente", "9", "No", "Parcial"],
    ["importacion", "ETL de lotes (novedades / pagos)", "8", "No", "Parcial"],
    ["comunicacion", "Mensajes y notificaciones", "6", "Sí", "Fina"],
    ["auditoria", "Incidencias + listas", "4", "Sí", "Fina"],
    ["contaduria", "Placeholder (fuera del dominio Rentas)", "1", "No", "Vacío"],
], widths=[1.4, 3.2, 0.8, 0.8, 0.8])

# ── 4. Comunicación ──────────────────────────────────────────────────
h1("4. Cómo se comunican los módulos")
bullet("Autenticación (todos → seguridad): cada request autenticado dispara una validación del "
       "token contra seguridad (GET /auth/me), implementada en la dependencia compartida "
       "shared/base_module.create_auth_dependency.")
bullet("Padrón (emisiones → ingresos_publicos): el paso 2 de la emisión (\"Cargar padrón\") trae "
       "por HTTP los inmuebles/comercios/vehículos con su base imponible desde ingresos_publicos.")
bullet("Auditoría (middleware en los 11 módulos): registra método, ruta, status y usuario, pero a "
       "los LOGS de la aplicación (stdout) — NO persiste en el módulo auditoria.")
bullet("Orquestación desde el frontend: el frontend combina datos de varios módulos sin que estos "
       "se llamen entre sí (ej. Vista 360 = ingresos_publicos para datos + emisiones para deuda).")
bullet("Redis: administracion, comunicacion, emisiones, tesoreria y seguridad lo declaran para "
       "workers asíncronos / cache.")

# ── 5. Matriz de interacción ─────────────────────────────────────────
h1("5. Matriz de interacción módulo a módulo")
table(["Llamador", "Llamado", "Vía", "Propósito", "Cuándo"], [
    ["Todos los módulos", "seguridad", "HTTP GET /auth/me", "Validar el JWT y traer datos/permisos del usuario", "Cada request autenticado"],
    ["emisiones", "ingresos_publicos", "HTTP (padrón)", "Traer el padrón (base imponible) para liquidar", "Paso 2 de cada emisión"],
    ["Frontend", "ingresos_publicos + emisiones", "HTTP (nginx)", "Vista 360: datos del contribuyente + deuda", "Por consulta"],
    ["Middleware (todos)", "— (stdout)", "Log", "Rastro de auditoría (no llega al módulo auditoria)", "Cada request"],
], widths=[1.5, 1.7, 1.2, 2.2, 1.2])
para("No existen otros llamados de negocio módulo-a-módulo: el resto de los módulos son "
     "autónomos y solo dependen de seguridad para autenticar.", italic=True, color=GREY)

# ── 6. Duplicación / sub-utilización ─────────────────────────────────
h1("6. Módulos sub-utilizados y duplicación de información")
para("Del mapeo anterior surgen estos puntos de atención:", bold=True)
table(["Tema", "Qué pasa", "Riesgo", "Recomendación"], [
    ["Personas duplicadas", "administracion_personas_fisicas/juridicas conviven con ingresos_publicos_personas/contribuyentes", "Datos de persona en 2 módulos → inconsistencia / doble carga", "Fuente única: maestro en administracion; ingresos_publicos referencia por HTTP"],
    ["Auditoría desconectada", "El middleware audita a stdout; el módulo auditoria (incidencias) no recibe ese rastro", "\"Auditoría\" en 2 lugares sin integración; rastro no consultable", "Cablear el middleware para POST no bloqueante al módulo auditoria"],
    ["wav vs ingresos_publicos", "wav tiene cuentas/DD.JJ./planes/pagos propios que solapan con ingresos_publicos", "Duplicidad de cuentas y DD.JJ. del contribuyente", "Consolidar en ingresos_publicos o delimitar wav como tributo con referencia"],
    ["interface vs comunicacion", "Ambos exponen notificaciones", "Duplicidad del canal de notificaciones", "Unificar notificaciones en comunicacion"],
    ["contaduria", "1 endpoint placeholder, sin uso; ya fuera del frontend", "Módulo muerto", "Eliminar o definir alcance real"],
    ["importacion / interface", "Backend sin pantalla en la app", "No se usan hoy", "Construir UI o postergar según prioridad"],
    ["Recaudación ↔ deuda", "tesoreria (cobro) y emisiones (cuenta corriente) no se comunican; el pago que baja la deuda vive dentro de emisiones", "Cobro real desacoplado de la deuda → riesgo de doble registro", "Integrar: la recaudación de tesoreria debe impactar la cuenta corriente de emisiones por HTTP"],
    ["Padrón (snapshot)", "emisiones_padron_contribuyentes copia el padrón para liquidar", "Copia por diseño (foto del momento)", "Aceptable; documentarlo como snapshot inmutable de la emisión"],
], widths=[1.2, 2.3, 1.7, 2.2])

# ── 7. Conclusiones ──────────────────────────────────────────────────
h1("7. Conclusiones")
bullet("La arquitectura de microservicios está bien planteada: bajo acoplamiento, HTTP-only, "
       "seguridad centralizada y auditoría no bloqueante.")
bullet("El dominio de Rentas está concentrado en 4 módulos maduros (ingresos_publicos, emisiones, "
       "tesoreria, administracion) + seguridad.")
bullet("Los focos a resolver son de INTEGRACIÓN y DUPLICACIÓN, no de arquitectura: unificar "
       "personas, integrar recaudación con la deuda, conectar la auditoría real y decidir el "
       "futuro de wav, interface, importacion y contaduria.")
note("Este documento refleja el estado del código al momento de generarse (conteo de endpoints, "
     "URLs de configuración y llamadas HTTP reales). La madurez se evaluó por estructura, no por "
     "pruebas funcionales exhaustivas.")

out = os.path.join(HERE, "Cheyenne-Arquitectura-y-Modulos.docx")
doc.save(out)
print("OK ->", out)
