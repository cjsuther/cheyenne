#!/usr/bin/env python3
"""Genera el Word de Análisis COMPLETO y Diseño del super-módulo Presupuesto de Cheyenne (v2)."""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(__file__)
PRIMARY = RGBColor(0x1D, 0x4E, 0xD8)
DARK = RGBColor(0x1F, 0x29, 0x37)
GREY = RGBColor(0x6B, 0x72, 0x80)

doc = Document()
n = doc.styles["Normal"]; n.font.name = "Calibri"; n.font.size = Pt(11)
n.paragraph_format.space_after = Pt(6)


def shade(cell, hexc):
    tcPr = cell._tc.get_or_add_tcPr(); sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear"); sh.set(qn("w:color"), "auto"); sh.set(qn("w:fill"), hexc)
    tcPr.append(sh)


def h1(t):
    p = doc.add_heading(t, level=1)
    for r in p.runs: r.font.color.rgb = PRIMARY
    return p


def h2(t):
    p = doc.add_heading(t, level=2)
    for r in p.runs: r.font.color.rgb = DARK
    return p


def para(t="", bold=False, italic=False, color=None):
    p = doc.add_paragraph(); r = p.add_run(t)
    r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = color
    return p


def bullet(t, lvl=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + 0.2 * lvl)
    p.add_run(t); return p


def code(t):
    p = doc.add_paragraph(); r = p.add_run(t)
    r.font.name = "Consolas"; r.font.size = Pt(8.5)
    pPr = p._p.get_or_add_pPr(); sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear"); sh.set(qn("w:color"), "auto"); sh.set(qn("w:fill"), "F1F5F9")
    pPr.append(sh)
    return p


def table(headers, rows, widths=None, fs=9):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
    for i, htext in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""
        run = c.paragraphs[0].add_run(htext); run.bold = True; run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); shade(c, "1D4ED8")
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""; run = cells[i].paragraphs[0].add_run(str(v)); run.font.size = Pt(fs)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    return t


def note(t):
    p = doc.add_paragraph(); r = p.add_run("  ★  " + t); r.italic = True; r.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)
    pPr = p._p.get_or_add_pPr(); sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear"); sh.set(qn("w:color"), "auto"); sh.set(qn("w:fill"), "EFF6FF"); pPr.append(sh)
    return p


# ══ Portada ══════════════════════════════════════════════════════════
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Cheyenne — Super-módulo Presupuesto"); r.bold = True; r.font.size = Pt(27); r.font.color.rgb = PRIMARY
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run("Análisis completo y diseño (v2)"); r.font.size = Pt(16); r.font.color.rgb = DARK
s2 = doc.add_paragraph(); s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s2.add_run("Basado en la extracción campo a campo del legacy vbPresupuesto (clases VB6 + DDL SQL reales) · "
               "documento de aprobación previa al desarrollo")
r.italic = True; r.font.color.rgb = GREY
doc.add_paragraph()

# ══ 1. Resumen ═══════════════════════════════════════════════════════
h1("1. Resumen ejecutivo")
para("Presupuesto será un SUPER-MÓDULO de Cheyenne: aparece como módulo propio en el dashboard "
     "posterior al login (junto a Seguridad, Contaduría, Ingresos Públicos, Contribuyente y "
     "Configuración) con su propio submenú. Técnicamente es un microservicio autónomo "
     "(modules/presupuesto) desplegable e implementable por sí solo, con seguridad integrada al "
     "módulo seguridad y permisos por acción de negocio.")
para("El análisis se completó con tres extracciones profundas del código fuente legacy (vbPresupuesto, "
     "vbComun, DDL en ScriptsSQL): diccionario de datos real, reglas de negocio con cita de archivo, "
     "inventario F1-F11 y mapa de consultas/reportes. Hallazgo central: el legacy YA es un sistema de "
     "movimientos (tabla Movimientos + agregación AuxPresupuesto); el diseño propuesto formaliza ese "
     "ledger y lo hace en tiempo real, con workflow de aprobación y API de afectación que el legacy "
     "no tiene.")

# ══ 2. Super-módulo ══════════════════════════════════════════════════
h1("2. Presupuesto como super-módulo")
bullet("Dashboard: nueva tarjeta 'Presupuesto' (descripción: 'Formulación, modificaciones y ejecución'). El dashboard queda con 6 super-módulos.")
bullet("Submenú propio (sidebar contextual): Partidas · Modificaciones · Recursos · Cuotas (fase 2) · Tablero · Nomencladores.")
bullet("Registro en config/superModules.js y en el buscador de secciones (navigation.js) con deep-links por solapa.")
bullet("Visibilidad por permiso: la tarjeta y los ítems se muestran solo si el rol tiene presupuesto_read (patrón existente).")

# ══ 3. Análisis del legacy ═══════════════════════════════════════════
h1("3. Análisis del legacy (relevado campo a campo)")
h2("3.1 La partida presupuestaria (PresupuestoDeGastos)")
para("Clave única de 5 dimensiones + crédito anual:")
code("(Ejer_Ejercicio, juri_Agrupamiento, espr_Agrupamiento, fufi_Agrupamiento, godg_Agrupamiento)\n"
     "→ prga_CostoAnual decimal(18,2) · prga_Fecha · auditoría (user_id, audi_Fecha, audi_Accion)\n"
     "Ejemplo real: 2010 · 1.1.1.01.01.000 (Adm. Central) · 01 (Conducción) · 1.1.0 (Tesoro) · 1.1.1 (Retribuciones)")
h2("3.2 Nomencladores (jerárquicos, con nivel y baja lógica)")
table(["Nomenclador", "Código", "Jerarquía", "Particularidades"], [
    ["Jurisdicciones (juri_)", "varchar(64) punteado '1.1.1.01.01.000'", "Nivel 1..4+ (juri_Nivel)", "Versionado por ejercicio; tipo (A/U)"],
    ["Objetos del gasto (godg_)", "'1.1.1' inciso.principal.parcial", "godg_Nivel", "Flags: grupo ocupacional, operación/consumo/producción, relación con finalidad-función"],
    ["Fuentes de financ. (fufi_)", "'1.1.0'", "fufi_Nivel", "Versionado por ejercicio"],
    ["Estructura programática (espr_)", "'01', '17' (plano)", "Por jurisdicción (la clave incluye juri)", "Tipo de programa, oficina responsable, textos descriptivos, PPG (género) y % PPG, código de préstamo"],
    ["Rubros de recursos (repr_)", "'1.1.4.02'", "Nivel 2..4", "Carácter económico obligatorio (desde 2011); mapeo contable en MatrizRecursosPorRubro"],
], widths=[1.7, 1.7, 1.4, 2.1], fs=8.5)
h2("3.3 El ledger del legacy (Movimientos + AuxPresupuesto)")
bullet("Tabla Movimientos: cada hecho presupuestario es una fila con la clave 5-dim, el comprobante origen (tipo/oficina/número/renglón), la columna (Tcol_Columna) y el importe con signo.")
bullet("Columnas: 1 saldo anterior · 2 ajustes · 4 ingresos (crédito inicial) · 5 modificaciones/pendiente preventivo · 6 compromiso · 11 devengado.")
bullet("AuxPresupuesto materializa los saldos por partida/mes: disponible = (1)+(2)+(4)−(5)−(6)−(11).")
bullet("Comprobantes de modificación: tipo 71 = ampliación, 72 = reducción; los items son filas de Movimientos con Tcol=5.")
h2("3.4 Reglas de las modificaciones (ModificacionesPresupuestarias.frm, 5.069 líneas)")
bullet("Obligatorios por ítem: las 4 dimensiones + detalle + importe ≠ 0 (líneas 4135-4162).")
bullet("En modo RAFAM: una reducción valida el saldo disponible de la partida — NO se permite sobregiro (línea 4159).")
bullet("No crea partidas: solo modifica crédito de partidas existentes.")
bullet("NO exige balance cero global (ampliación y reducción son actos separados); requiere acto administrativo (tipo + número de decreto/resolución) y admite expediente.")
bullet("Anulación: invierte los signos de la modificación original (restituye el crédito); numeración automática opcional.")
h2("3.5 Recursos")
bullet("CalculoDeRecursos: PK (ejercicio, jurisdicción, rubro) → estimado / programado / ejecutado + descripción metodológica.")
bullet("El 'programado real' se deriva de MovimientosRecursos cuando no hay valor directo — otro ledger embrionario.")
bullet("MatrizRecursosPorRubro: mapeo contable (cuentas debe/haber para devengado y percibido) — pertenece a la integración con Contaduría, no al MVP.")
bullet("Modificaciones al cálculo (ModificCalcRec.frm): con acto administrativo; sin balanceo automático contra gastos.")
h2("3.6 Cuotas de compromiso")
bullet("Dos niveles: ProgramacionDelCompromiso (cuota AUTORIZADA por trimestre, clave 5-dim + trimestre) y CuotasDelCompromiso (cuota SOLICITADA por categoría contable).")
bullet("El compromiso en Contaduría valida contra la cuota del trimestre (control de caja del gasto).")
h2("3.7 Formulación F1-F11 y evaluación — veredicto de uso")
table(["Formulario", "Captura", "Veredicto"], [
    ["F1 Política presupuestaria", "Textos: situación actual, políticas, incidencia futura", "Documental — posponer"],
    ["F2 Programación de recursos", "Mensual/trimestral por rubro", "Central — cubierto por Recursos (simplificado)"],
    ["F3/F4 Estructura programática", "Alta del árbol programático + descripciones", "Central — es el nomenclador"],
    ["F5 Metas por programa", "Metas físicas con unidad de medida, por trimestre", "Importante — fase 3 (tablero primero)"],
    ["F6/F72 RRHH por cargo", "Cargos, categorías, cantidades, hs cátedra, costo anual", "Importante (~60% del gasto) — fase 2"],
    ["F7 Bienes y servicios", "Cantidad × precio por objeto, importa Excel", "Útil — cubierto por carga masiva de partidas"],
    ["F8/F9 Proyectos físico/financiero", "Avance físico/financiero plurianual", "Sectorial — fase 3"],
    ["F10 Deuda / F11 Transferencias", "Préstamos y transferencias presupuestadas", "Sectorial — fase 3 (o módulo Crédito Público)"],
    ["Evaluación de desvíos", "% desvío + causa + conclusiones por trimestre", "Reemplazado por tablero de ejecución"],
], widths=[1.8, 2.7, 2.4], fs=8.5)
note("Decisión de alcance: el MVP implementa el corazón operativo (partidas, modificaciones, recursos, "
     "ejecución/afectaciones, tablero). RRHH por cargo y cuotas van en fase 2; metas y proyectos en "
     "fase 3. Nada se descarta sin registro: queda etiquetado en este documento.")

# ══ 4. Decisiones de diseño ══════════════════════════════════════════
h1("4. Decisiones de diseño (DD)")
table(["#", "Decisión", "Fundamento"], [
    ["DD-01", "Ledger como fuente única de verdad: presupuesto_movimientos; los saldos SIEMPRE se derivan, nunca se almacenan", "Formaliza lo que el legacy hace a medias (Movimientos + AuxPresupuesto batch); elimina inconsistencias"],
    ["DD-02", "Nomencladores GLOBALES con vigencia (activo/fecha_baja), no versionados por ejercicio", "El legacy copia nomencladores por año (costo de administración alto); la partida ya fija el ejercicio; simplifica la prórroga"],
    ["DD-03", "Estructura programática asociada a jurisdicción (opcional)", "Fiel al legacy (espr pertenece a juri); si el municipio usa estructura única, se deja juri nula"],
    ["DD-04", "Modificación con 3 tipos: ampliación (+), reducción (−), compensación (Σ=0 validado)", "Supera al legacy (71/72 sin balance); la compensación es el caso más común y el más propenso a error manual"],
    ["DD-05", "Control de sobregiro configurable por ejercicio (bloquear/advertir)", "Replica la regla RAFAM de no-sobregiro como default, con flexibilidad"],
    ["DD-06", "Cadena de afectación: preventivo → compromiso → devengado → pagado; cada etapa puede referenciar y cerrar la anterior", "Semántica limpia de saldos: disponible = vigente − preventivos vigentes − comprometido"],
    ["DD-07", "API de afectación idempotente por (origen, referencia, tipo)", "Contrato para Contaduría (futuro ciclo del gasto) sin doble registro ante reintentos"],
    ["DD-08", "Todos los códigos jerárquicos como texto punteado + id_padre + nivel derivado", "Compatible con la codificación RAFAM del legacy y navegable como árbol"],
    ["DD-09", "Sin metas/PPG/proyectos en el núcleo: etiquetas JSON libres por partida", "Cubre clasificaciones transversales (p.ej. género) sin costo estructural"],
    ["DD-10", "Migración inicial de datos desde el legacy soportada de fábrica (importadores)", "Existen DDL y datos (ScriptsSQL); un municipio arranca cargando su presupuesto real"],
], widths=[0.6, 3.2, 3.1], fs=8.5)

# ══ 5. Modelo de datos ═══════════════════════════════════════════════
h1("5. Diccionario de datos (Cheyenne)")
para("Convenciones: BigInteger id PK autoincremental, activo bool (soft-delete), created_at/updated_at, "
     "paginación skip/limit y ordenamiento sort_by/sort_dir en todos los listados (infra compartida).")
h2("5.1 Nomencladores")
table(["Tabla", "Columnas propias", "Trazabilidad legacy"], [
    ["presupuesto_jurisdicciones", "codigo (punteado, unique), nombre, tipo, id_padre, nivel (derivado), activo", "Jurisdicciones (juri_*)"],
    ["presupuesto_estructuras", "codigo, nombre, tipo_programa, id_jurisdiccion (nullable), descripcion, activo — unique(codigo, id_jurisdiccion)", "EstructuraProgramatica (espr_*)"],
    ["presupuesto_objetos_gasto", "codigo (inciso.principal.parcial, unique), nombre, detalle, id_padre, nivel, activo", "GastosObjetoDelGasto (godg_*)"],
    ["presupuesto_fuentes", "codigo (unique), nombre, detalle, origen, id_padre, nivel, activo", "FuenteFinanciamiento (fufi_*)"],
    ["presupuesto_rubros", "codigo (unique), nombre, id_padre, nivel, caracter_economico, activo", "RecursosPorRubro (repr_*)"],
], widths=[1.9, 3.4, 1.6], fs=8.5)
h2("5.2 Núcleo")
table(["Tabla", "Columnas propias", "Notas"], [
    ["presupuesto_ejercicios", "anio (unique), estado, control_sobregiro (bloquear/advertir), acto_administrativo, fecha_aprobacion, fechas/usuarios por transición", "Workflow §7.1"],
    ["presupuesto_partidas", "anio, id_jurisdiccion, id_estructura, id_objeto_gasto, id_fuente, credito_inicial numeric(18,2), descripcion, etiquetas JSON — unique(5 dimensiones)", "≡ PresupuestoDeGastos"],
    ["presupuesto_movimientos", "id_partida, fecha, tipo (inicial|modificacion|preventivo|compromiso|devengado|pagado|liberacion|ajuste), importe (signo), origen (modulo), referencia_tipo, referencia, id_afectacion (nullable), id_usuario, usuario_nombre, observaciones", "LEDGER ≡ Movimientos+Tcol"],
    ["presupuesto_modificaciones", "numero (auto por ejercicio), anio, fecha, tipo (ampliacion|reduccion|compensacion), acto_administrativo (oblig.), expediente, estado (borrador|aprobada|anulada), observaciones, usuario/fecha de alta, aprobación y anulación", "≡ Comprobantes 71/72 + workflow"],
    ["presupuesto_modificacion_items", "id_modificacion, id_partida, importe (signo), detalle (oblig.)", "≡ Movimientos Tcol=5"],
    ["presupuesto_afectaciones", "id_partida, tipo (preventivo|compromiso|devengado|pagado), importe, estado (vigente|cerrada|liberada), id_origen (afectación previa de la cadena), origen_modulo, referencia_tipo, referencia — unique(origen_modulo, referencia_tipo, referencia, tipo)", "Cadena DD-06; genera movimientos"],
    ["presupuesto_recursos", "anio, id_jurisdiccion, id_rubro, estimado_inicial, metodologia — unique(anio, juri, rubro)", "≡ CalculoDeRecursos"],
    ["presupuesto_recurso_movimientos", "id_recurso, tipo (inicial|modificacion|percibido), importe, referencia, usuario", "ledger del recurso"],
    ["presupuesto_cuotas (fase 2)", "anio, id_jurisdiccion, id_fuente (null), id_objeto_gasto nivel inciso (null), trimestre, importe_autorizado", "≡ ProgramacionDelCompromiso"],
], widths=[1.8, 3.7, 1.4], fs=8.5)
h2("5.3 Saldos derivados (endpoint, no tabla)")
code("credito_vigente = credito_inicial + Σ modificaciones aprobadas (items de la partida)\n"
     "preventivo      = Σ afectaciones preventivo VIGENTES (no convertidas ni liberadas)\n"
     "comprometido    = Σ afectaciones compromiso (vigentes o devengadas)\n"
     "devengado       = Σ afectaciones devengado · pagado = Σ pagado\n"
     "disponible      = credito_vigente − preventivo − comprometido\n"
     "(equivale a la fórmula del legacy (1)+(2)+(4)−(5)−(6) con semántica de cadena explícita)")

# ══ 6. Reglas de negocio ═════════════════════════════════════════════
h1("6. Reglas de negocio (RN)")
table(["#", "Regla"], [
    ["RN-01", "La partida es única por (ejercicio, jurisdicción, estructura, objeto del gasto, fuente). El alta valida nomencladores activos."],
    ["RN-02", "Partidas: alta/edición/carga masiva SOLO con el ejercicio en formulación. Desde la aprobación, el crédito solo cambia por modificación aprobada."],
    ["RN-03", "Modificación: cada ítem exige partida existente, detalle e importe ≠ 0 (regla del legacy). No crea partidas."],
    ["RN-04", "Compensación: Σ importes = 0 (validado al guardar y al aprobar). Ampliación: todos > 0. Reducción: todos < 0."],
    ["RN-05", "Al aprobar una reducción (o compensación con ítems negativos): valida disponible por partida según control_sobregiro del ejercicio (bloquear = rechaza; advertir = permite y registra)."],
    ["RN-06", "Aprobación de modificación: genera un movimiento por ítem; anulación: genera los contra-movimientos. Nada se edita ni borra después de aprobada."],
    ["RN-07", "Acto administrativo obligatorio en toda modificación (tipo + número); expediente opcional."],
    ["RN-08", "Numeración de modificaciones: automática, secuencial por ejercicio, sin huecos visibles."],
    ["RN-09", "Afectación: valida ejercicio vigente y disponible (según control_sobregiro). Idempotente por (origen, referencia, tipo)."],
    ["RN-10", "Cadena: un compromiso puede referenciar un preventivo (lo cierra y libera su reserva); un devengado a un compromiso; un pagado a un devengado. También se admiten etapas directas (compromiso sin preventivo)."],
    ["RN-11", "Liberación: solo sobre afectaciones vigentes; genera contra-movimiento; deja rastro de usuario/fecha/motivo."],
    ["RN-12", "Prórroga de ejercicio: copia partidas (con crédito inicial = vigente del año anterior u original, a elección) al nuevo año en estado formulación. No copia movimientos."],
    ["RN-13", "Cierre de ejercicio: bloquea toda nueva afectación/modificación; requiere que no queden modificaciones en borrador."],
    ["RN-14", "Recursos: estimado inicial editable en formulación; luego solo por modificación de recursos (mismo workflow simple). El percibido llega por API (Tesorería/Rentas, futuro)."],
    ["RN-15", "Todos los importes: numeric(18,2), redondeo half-up, moneda única (pesos)."],
], widths=[0.6, 6.3], fs=8.7)

# ══ 7. Workflows y permisos ══════════════════════════════════════════
h1("7. Workflows y seguridad")
h2("7.1 Ejercicio")
code("FORMULACION → (aprobar: acto admin.) → APROBADO → (poner en vigencia) → VIGENTE → (cerrar) → CERRADO\n"
     "Transiciones con permiso propio, usuario y fecha registrados. Reapertura: solo admin, quedando auditada.")
h2("7.2 Modificación presupuestaria")
code("BORRADOR → (aprobar) → APROBADA → (anular) → ANULADA\n"
     "Borrador editable; aprobada intocable (solo anulable con contra-movimientos).")
h2("7.3 Matriz de permisos (seed en seguridad por migración)")
table(["Permiso", "Habilita", "Rol típico"], [
    ["presupuesto_read", "Ver el super-módulo, partidas, saldos, tablero", "Consulta"],
    ["presupuesto_write", "Nomencladores, partidas en formulación, borradores de modificación, recursos", "Analista de presupuesto"],
    ["presupuesto_delete", "Bajas lógicas de nomencladores/partidas en formulación", "Analista senior"],
    ["presupuesto_modificacion_aprobar", "Aprobar/anular modificaciones (gastos y recursos)", "Director de Presupuesto"],
    ["presupuesto_aprobar_ejercicio", "Aprobar, poner en vigencia, cerrar y prorrogar ejercicios", "Secretario de Hacienda"],
    ["presupuesto_afectar", "Registrar/liberar afectaciones vía API o pantalla", "Contaduría (servicio) / operador autorizado"],
    ["presupuesto_admin", "Todo lo anterior + configuración (control_sobregiro, reapertura)", "Administrador"],
], widths=[2.1, 3.3, 1.5], fs=8.5)

# ══ 8. Especificación funcional por pantalla ═════════════════════════
h1("8. Especificación funcional (pantallas del super-módulo)")
h2("8.1 Partidas (pantalla central)")
bullet("Cabecera: selector de ejercicio con badge de estado + acciones del ejercicio según permiso (aprobar/vigencia/cerrar/prorrogar).")
bullet("Grilla: dimensiones como 'nombre (código)', crédito inicial, modificaciones, vigente, preventivo, comprometido, devengado, pagado, DISPONIBLE (verde/rojo). Filtro por columna, ordenamiento, paginación.")
bullet("Fila → drawer con el ledger de la partida (movimientos con origen, referencia, usuario, fecha).")
bullet("Acciones: Nueva partida (solo formulación; validación RN-01), Importar planilla (CSV con vista previa de errores por fila), Exportar CSV.")
h2("8.2 Modificaciones")
bullet("Listado con número, fecha, tipo, acto administrativo, estado (badge), total ampliado/reducido, usuario.")
bullet("Asistente 3 pasos: (1) tipo + acto administrativo + expediente + observaciones; (2) items — buscador de partida + importe + detalle, con Σ en vivo y semáforo de balance (compensación) y de disponible (reducción); (3) resumen → guarda borrador.")
bullet("Detalle: items + historial (creada/aprobada/anulada por quién y cuándo) + botones Aprobar/Anular según permiso (RN-05/06).")
h2("8.3 Recursos")
bullet("Grilla por jurisdicción × rubro (árbol de rubros): estimado inicial, modificaciones, vigente, percibido. Alta/edición en formulación; modificaciones con el mismo asistente (tipo único).")
h2("8.4 Tablero")
bullet("Tarjetas del ejercicio: vigente, comprometido, devengado, pagado, disponible, % ejecución.")
bullet("Barras por inciso (nivel 1 de objeto), por jurisdicción y por fuente; equilibrio gastos vs. recursos; últimas modificaciones y afectaciones.")
h2("8.5 Nomencladores")
bullet("Una solapa por nomenclador (CrudTab): código, nombre, padre (search), extras propios. Vista de árbol plegable para jerárquicos. Importadores: desde administracion (jurisdicciones) y desde CSV legacy.")
h2("8.6 Cuotas (fase 2)")
bullet("Matriz trimestre × dimensión (jurisdicción/fuente/inciso) de importes autorizados; validación de compromisos contra cuota (RN futura).")

# ══ 9. API ═══════════════════════════════════════════════════════════
h1("9. API completa")
code("GET /health\n"
     "# Nomencladores (CRUD estándar + filtros + sort): /jurisdicciones /estructuras /objetos-gasto /fuentes /rubros\n"
     "POST /jurisdicciones/importar        # opcional, desde administracion\n"
     "POST /nomencladores/importar-legacy  # CSV del legacy (jurisdicciones/objetos/fuentes/estructuras/rubros)\n"
     "# Ejercicios\n"
     "GET/POST /ejercicios · GET /ejercicios/{anio}\n"
     "POST /ejercicios/{anio}/aprobar | /vigencia | /cerrar | /reabrir | /prorrogar {origen_credito: inicial|vigente}\n"
     "# Partidas\n"
     "GET  /partidas?anio&…dims&sort_by&sort_dir&skip&limit   → filas con saldos §5.3\n"
     "POST /partidas · PUT/DELETE /partidas/{id}              (RN-02)\n"
     "POST /partidas/importar   (CSV, dry_run=true para validar)\n"
     "GET  /partidas/{id} · GET /partidas/{id}/movimientos\n"
     "GET  /resumen?anio&por=inciso|jurisdiccion|fuente · GET /export?anio\n"
     "# Modificaciones (gastos)\n"
     "GET/POST /modificaciones · GET/PUT/DELETE /modificaciones/{id}   (PUT/DELETE solo borrador)\n"
     "POST /modificaciones/{id}/aprobar · /anular\n"
     "# Recursos\n"
     "GET/POST /recursos · PUT /recursos/{id}\n"
     "GET/POST /recursos-modificaciones · POST /recursos-modificaciones/{id}/aprobar | /anular\n"
     "POST /recursos/percibido   {id_recurso|dimensiones, importe, referencia}   # futuro: Tesorería/Rentas\n"
     "# Afectaciones (contrato del ciclo del gasto)\n"
     "GET/POST /afectaciones · POST /afectaciones/{id}/liberar {motivo}\n"
     "  POST body: { tipo, id_partida | dimensiones{...}, importe, id_origen?, origen_modulo,\n"
     "               referencia_tipo, referencia, observaciones }  → 409 si excede disponible (bloquear)\n"
     "# Cuotas (fase 2): GET/POST /cuotas · validación en afectación tipo compromiso")
note("Errores consistentes: 400 validación (detalle por campo), 403 permiso, 404 inexistente, 409 regla "
     "de negocio (sobregiro, balance, estado). Todos los listados usan la infra compartida de filtros y orden.")

# ══ 10. Migración desde el legacy ════════════════════════════════════
h1("10. Migración de datos desde el legacy")
bullet("Fuentes disponibles: DDL y datos en legacy/ScriptsSQL (dbo_PresupuestoDeGastos, dbo_Jurisdicciones, dbo_GastosObjetoDelGasto, dbo_FuenteFinanciamiento, dbo_EstructuraProgramatica, dbo_RecursosPorRubro, dbo_CalculoDeRecursos).")
bullet("Orden: nomencladores (mapa código legacy → id Cheyenne) → partidas del ejercicio elegido (prga_CostoAnual → credito_inicial) → recursos (estimado). Los movimientos históricos NO se migran (arranque limpio con saldo inicial).")
bullet("Importadores con dry-run: reporte de filas aceptadas/rechazadas antes de confirmar.")
bullet("Verificación: totales por inciso y por jurisdicción comparados contra los reportes del legacy (PresupuestoDeGastos.rpt).")

# ══ 11. Reportes ═════════════════════════════════════════════════════
h1("11. Reportes y salidas")
bullet("MVP: exportación CSV de partidas/saldos y resumen por inciso/jurisdicción/fuente (equivalente funcional de PresupuestoDeGastos.rpt, PresupuestoPorObjetoDelGasto.rpt, PresupuestoFuenteFinanciamiento.rpt, ModificacionesPresupuestarias.rpt).")
bullet("Fase 2: PDF con reportlab (patrón de Emisiones) para los formatos de presentación (F12 Resumen por inciso, Cuenta Ahorro-Inversión).")
bullet("Formatos RAFAM oficiales: se abordan con la integración a Contaduría (rendición), no en el MVP.")

# ══ 12. NFR ══════════════════════════════════════════════════════════
h1("12. Requisitos no funcionales")
bullet("Consulta de partidas con saldos < 1,5 s para 10.000 partidas (agregación SQL sobre el ledger con índices por id_partida+tipo; vista materializada opcional si crece).")
bullet("Paginación obligatoria (max 100), ordenamiento server-side, filtros por columna (infra compartida).")
bullet("Auditoría: middleware compartido (rastro de accesos) + usuario/fecha en cada transición y movimiento.")
bullet("Transaccionalidad: aprobar modificación / registrar afectación = una transacción DB (todo o nada).")
bullet("Concurrencia: validación de disponible con bloqueo de fila de la partida (SELECT FOR UPDATE) al afectar.")
bullet("Despliegue: contenedor propio + ruta Nginx /api/presupuesto/ + migración SQL idempotente + seed de permisos.")

# ══ 13. Plan de entregas ═════════════════════════════════════════════
h1("13. Plan de entregas y criterios de aceptación")
table(["Entrega", "Contenido", "Criterios de aceptación"], [
    ["E1 Esqueleto + super-módulo", "Servicio, compose, nginx, health, auth, permisos seed; tarjeta en dashboard + submenú + buscador", "/api/presupuesto/health OK; tarjeta visible solo con presupuesto_read; permisos asignables en Seguridad"],
    ["E2 Nomencladores + Ejercicios", "5 CRUDs jerárquicos + workflow ejercicio + importadores", "Árbol de objetos del gasto navegable; ejercicio 2027 creado en formulación; jurisdicciones importadas"],
    ["E3 Partidas + Ledger", "Partidas, carga masiva dry-run, movimientos, saldos derivados, export, drawer de ledger", "Partida con inicial 100 muestra disponible 100; import CSV con 1 error rechaza solo esa fila; RN-01/02 verificadas"],
    ["E4 Modificaciones", "Asistente, workflow, balance Σ=0, no-sobregiro, anulación con contra-mov.", "Compensación desbalanceada no aprueba; reducción mayor al disponible rechazada (modo bloquear); anulación restituye vigente"],
    ["E5 Afectaciones + Recursos + Tablero", "API afectación con cadena e idempotencia, liberaciones; recursos + modificaciones; tablero", "Preventivo 100 baja disponible; compromiso sobre ese preventivo no double-cuenta; retry con misma referencia no duplica; tablero cuadra con la grilla"],
], widths=[1.5, 2.9, 2.5], fs=8.5)
para("Cada entrega se despliega y verifica end-to-end en el servidor (patrón actual: build + migración + "
     "verificación por API + limpieza de datos de prueba).", italic=True, color=GREY)

# ══ 14. Preguntas abiertas ═══════════════════════════════════════════
h1("14. Preguntas abiertas (a validar antes o durante E1)")
bullet("¿Control de sobregiro por defecto en 'bloquear' (recomendado, fiel a RAFAM) o 'advertir'?")
bullet("¿La estructura programática se usará por jurisdicción (como el legacy) o única para todo el municipio?")
bullet("¿Prórroga con crédito inicial = 'inicial del año anterior' o 'vigente del año anterior'? (el sistema soporta ambas; definir default)")
bullet("¿Se migran datos reales del legacy en la puesta en marcha (qué ejercicio)?")
bullet("¿PPG (perspectiva de género) como etiqueta libre alcanza, o se necesita el % formal por estructura (espr_PorcentajePPG)?")

out = os.path.join(HERE, "Cheyenne-Presupuesto-Analisis-y-Diseno.docx")
doc.save(out)
print("OK ->", out)
