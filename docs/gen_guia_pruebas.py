#!/usr/bin/env python3
"""Genera la guía de pruebas del flujo completo de Cheyenne en formato Word (.docx)."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PRIMARY = RGBColor(0x1D, 0x4E, 0xD8)   # azul
DARK = RGBColor(0x1F, 0x29, 0x37)
GREY = RGBColor(0x6B, 0x72, 0x80)

doc = Document()

# ---- estilos base ----
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)

def shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear'); sh.set(qn('w:color'), 'auto'); sh.set(qn('w:fill'), hex_color)
    tcPr.append(sh)

def h1(text):
    p = doc.add_heading(text, level=1)
    for r in p.runs: r.font.color.rgb = PRIMARY
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    for r in p.runs: r.font.color.rgb = DARK
    return p

def para(text='', bold=False, italic=False, color=None, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = color
    if size: r.font.size = Pt(size)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 + 0.25*level)
    p.add_run(text)
    return p

def step(n, text):
    p = doc.add_paragraph()
    r = p.add_run(f'{n}.  '); r.bold = True; r.font.color.rgb = PRIMARY
    p.add_run(text)
    return p

def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    for line in text.split('\n'):
        r = p.add_run(line + '\n')
        r.font.name = 'Consolas'; r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0x0B, 0x3D, 0x2E)
    shade_paragraph(p, 'F1F5F9')
    return p

def shade_paragraph(p, hex_color):
    pPr = p._p.get_or_add_pPr()
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear'); sh.set(qn('w:color'), 'auto'); sh.set(qn('w:fill'), hex_color)
    pPr.append(sh)

def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.style = 'Light Grid Accent 1'
    for i, htext in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ''
        run = c.paragraphs[0].add_run(htext); run.bold = True; run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade(c, '1D4ED8')
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            run = cells[i].paragraphs[0].add_run(str(val)); run.font.size = Pt(10)
    return t

def note(text):
    p = doc.add_paragraph()
    r = p.add_run('  ⚠  ' + text); r.italic = True; r.font.color.rgb = RGBColor(0x92, 0x40, 0x0E)
    shade_paragraph(p, 'FEF3C7')
    return p

# ════════════════════════════════════════════════════════════════════════
# PORTADA
# ════════════════════════════════════════════════════════════════════════
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('Cheyenne'); r.bold = True; r.font.size = Pt(34); r.font.color.rgb = PRIMARY
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run('Guía de Pruebas — Flujo Completo del Sistema'); r.font.size = Pt(15); r.font.color.rgb = DARK
s2 = doc.add_paragraph(); s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s2.add_run('Sistema de Ingresos Públicos · Motor de Cálculo de Rentas'); r.italic = True; r.font.color.rgb = GREY
doc.add_paragraph()

table(['Dato', 'Valor'], [
    ['URL del sistema', 'http://101.44.13.252/'],
    ['Usuario', 'admin'],
    ['Contraseña', 'Admin@2024!'],
    ['Navegador', 'Chrome / Edge / Firefox actualizado'],
    ['Importante', 'Si ves contenido viejo, recargá con Ctrl+Shift+R (Cmd+Shift+R en Mac)'],
])
doc.add_paragraph()
para('Esta guía recorre, de punta a punta, el circuito tributario: carga de la base imponible, '
     'emisión de un tributo con su motor de cálculo, generación de deuda y comprobantes, y planes de pago. '
     'Cada sección indica qué hacer, dónde hacerlo y qué resultado esperar.', italic=True, color=GREY)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
h1('1. Acceso al sistema')
step(1, 'Abrí http://101.44.13.252/ en el navegador.')
step(2, 'Ingresá usuario admin y contraseña Admin@2024! y presioná Iniciar sesión.')
step(3, 'Deberías ver el Dashboard con las tarjetas de los módulos.')
para('Resultado esperado:', bold=True)
bullet('El menú lateral izquierdo muestra: Dashboard, Administración, Ingresos Públicos, Emisiones, Tesorería, Comunicación, Auditoría, Seguridad.')
bullet('Arriba del menú hay una caja de búsqueda “Buscar sección…”.')
note('El menú solo muestra los módulos sobre los que el usuario tiene permiso. Con admin (superusuario) se ven todos.')

# ════════════════════════════════════════════════════════════════════════
h1('2. Buscador del menú')
para('Permite saltar directo a cualquier sección o solapa permitida.')
step(1, 'En la caja “Buscar sección…”, escribí: valuac')
step(2, 'En la lista que aparece, elegí “Valuaciones” (contexto: Ingresos Públicos).')
step(3, 'El sistema te lleva directo a Ingresos Públicos con la solapa Valuaciones abierta.')
para('Probá también:', bold=True)
bullet('Escribir “comprob”, “simular”, “cuotas”, “DD.JJ.” → cada uno salta a su solapa.')
bullet('Teclado: flechas ↑/↓ para moverte, Enter para entrar, Esc para limpiar.')
note('El buscador solo lista secciones a las que el usuario tiene acceso por permisos.')

# ════════════════════════════════════════════════════════════════════════
h1('3. Cargar la base imponible (inmuebles)')
para('La base imponible son los datos sobre los que el motor calcula el tributo. '
     'Para el ejemplo usaremos dos inmuebles.')

h2('3.1 Verificar inmuebles')
step(1, 'Ir a Ingresos Públicos › solapa Inmuebles.')
step(2, 'Confirmá que existan al menos 2 inmuebles (anotá sus IDs, p. ej. 1 y 2). Si no hay, creá con “+ Nuevo”.')

h2('3.2 Cargar Valuaciones')
step(1, 'Ir a la solapa Valuaciones › botón “+ Nuevo”.')
step(2, 'Cargá los siguientes dos registros:')
table(['Inmueble', 'Tipo Valuación', 'Ejercicio', 'Valor', 'Fecha Vigencia'], [
    ['#1', '1', '2026', '350000', '01/01/2026'],
    ['#2', '1', '2026', '200000', '01/01/2026'],
])

h2('3.3 Cargar Superficies')
step(1, 'Ir a la solapa Superficies › “+ Nuevo”.')
step(2, 'Cargá:')
table(['Inmueble', 'Tipo Superficie', 'Clase', 'Superficie (m²)', 'Fecha Vigencia'], [
    ['#1', '1', '1', '80', '01/01/2026'],
    ['#2', '1', '1', '50', '01/01/2026'],
])
para('Resultado esperado:', bold=True)
bullet('Cada alta aparece en la tabla; el inmueble se elige desde un combo que muestra “#id · nomenclatura”.')
note('La solapa Frentes es opcional para este ejemplo (la fórmula de demo no usa metros de frente).')

# ════════════════════════════════════════════════════════════════════════
h1('4. Definir la fórmula de cálculo (prerrequisito técnico)')
para('El motor calcula según las fórmulas de cada tasa (tabla FormulaTasa). Todavía no hay una pantalla '
     'para cargarlas, así que un administrador las carga una sola vez por SQL. La fórmula de demo del ABL es:')
code('a_cancelar = #I_VALUACION(0) * 0.012 + #I_SUPERFICIE(1,1) * 15\n'
     'a_pagar    = a_cancelar * 0.90        (10% de descuento por pago al 1er vencimiento)\n'
     'condición  = #I_VALUACION(0) > 0      (solo liquida cuentas con valuación)')
para('El SQL exacto para sembrar esta fórmula está en el Apéndice A.', italic=True, color=GREY)
note('En producción, esta fórmula se reemplaza por el catálogo real de tasas (dump de FormulaTasa). '
     'El SQL del Apéndice A es solo para probar el circuito.')

# ════════════════════════════════════════════════════════════════════════
h1('5. Emisión punta a punta (núcleo de la prueba)')
para('Una emisión recorre un workflow de 16 pasos. Vamos a generar las liquidaciones, la deuda en '
     'cuenta corriente y los comprobantes.')

h2('5.1 Crear la emisión')
step(1, 'Ir a Emisiones › botón “+ Nueva Emisión”.')
step(2, 'Completá: Tipo de tributo = inmuebles · Período = 2026 · Descripción = “Prueba ABL”.')
step(3, 'Guardar. La emisión aparece en la lista con estado inicial y Paso 0.')

h2('5.2 Ejecutar el workflow')
step(1, 'En la fila de la emisión, hacé clic en el botón de workflow (ícono de pasos).')
step(2, 'Se abre el modal con el Progreso (0/16). Ejecutá los pasos en orden con el botón “Ejecutar” '
        'que aparece en el paso actual:')
table(['Paso', 'Nombre', 'Qué hace / qué esperar'], [
    ['1', 'Validar Parámetros', 'Valida la emisión. Queda en verde.'],
    ['2', 'Cargar Padrón', 'Trae los inmuebles con su base imponible. Esperar “completado”.'],
    ['3', 'Validar Padrón', 'Total contribuyentes = 6 (o los que existan).'],
    ['4', 'Calcular Base Imponible', 'Marca cuántos tienen base. sin_base_imponible = 0.'],
    ['5–7', 'Alícuotas / Bonif. / Recargos', 'Informativos: el cálculo real ocurre en el paso 8.'],
    ['8', 'Generar Liquidaciones', '➜ Aparece el panel “Liquidaciones”. Ver tabla 5.3.'],
    ['9', 'Validar Liquidaciones', 'Paso de aprobación: tildar/observación y “Aprobar”.'],
    ['10', 'Generar Ordenamiento', 'Ordena las liquidaciones.'],
    ['11', 'Generar Cuentas Corrientes', '➜ Aparece el panel “Cuenta corriente”. Ver tabla 5.4.'],
    ['12', 'Generar Comprobantes', '➜ Aparece el panel “Comprobantes / Recibos”. Ver tabla 5.5.'],
    ['13–14', 'Imputación / Publicar Deuda', 'Pasos administrativos.'],
    ['15–16', 'Solicitar / Aprobar Emisión', 'Aprobación final (observación + “Aprobar”).'],
])

h2('5.3 Resultado esperado — Liquidaciones (tras el paso 8)')
para('El panel “Liquidaciones” debe mostrar 2 filas (solo los inmuebles con valuación):')
table(['Contrib.', 'Objeto', 'Tasa', 'Vto', 'a Cancelar', 'a Pagar'], [
    ['1', '1', '1', '1', '5400,00', '4860,00'],
    ['2', '2', '1', '1', '3150,00', '2835,00'],
])
para('Total a pagar (encabezado del panel): 7695,00', bold=True)
para('Verificación del cálculo (inmueble #1): 350000 × 0,012 + 80 × 15 = 4200 + 1200 = 5400 → '
     'con 10% de descuento = 4860.', italic=True, color=GREY)

h2('5.4 Resultado esperado — Cuenta corriente (tras el paso 11)')
para('El panel “Cuenta corriente” registra la deuda (importe a cancelar, sin descuento):')
table(['Contrib.', 'Concepto', 'Estado', 'Saldo'], [
    ['1', 'Emision inmuebles 2026 - tasa 1 vto 1', 'pendiente', '5400,00'],
    ['2', 'Emision inmuebles 2026 - tasa 1 vto 1', 'pendiente', '3150,00'],
])
para('Deuda total: 8550,00', bold=True)

h2('5.5 Resultado esperado — Comprobantes (tras el paso 12)')
para('El panel “Comprobantes / Recibos” muestra un recibo por cuenta, con código de barras:')
table(['Número', 'Contrib.', 'Código de barras', 'Importe'], [
    ['E000004-000001', '1', '00000400000001000000486000', '4860,00'],
    ['E000004-000002', '2', '00000400000002000000283500', '2835,00'],
])
para('El número de emisión (E000004) variará según el ID real de tu emisión.', italic=True, color=GREY)

# ════════════════════════════════════════════════════════════════════════
h1('6. Planes de pago')
h2('6.1 Simular un plan (sistema francés)')
step(1, 'Ir a Ingresos Públicos › solapa Simular Plan.')
step(2, 'Cargá: Monto total = 1000 · Cuotas = 3 · Interés mensual (%) = 5 · Anticipo = 0. Clic en “Simular”.')
para('Resultado esperado:', bold=True)
bullet('Tarjetas: Anticipo 0,00 · Financiado 1000,00 · Total intereses ≈ 101,63 · Total a pagar ≈ 1101,63.')
bullet('Tabla de cuotas: 3 cuotas de ≈ 367,21 con capital / interés / saldo decrecientes.')

h2('6.2 Generar y ver cuotas de un plan')
step(1, 'Ir a la solapa Cuotas Plan.')
step(2, 'Elegí un plan del combo y clic en “Generar cuotas”.')
step(3, 'Aparece el mensaje “✓ N cuotas generadas” y la tabla con capital, interés, importe y vencimiento.')

# ════════════════════════════════════════════════════════════════════════
h1('7. Base imponible de comercio y vehículo')
bullet('Ingresos Públicos › Com. Rubros: alta de rubros del comercio (combo de comercio + rubro).')
bullet('Ingresos Públicos › Com. DD.JJ.: declaraciones juradas de ingresos (período, mes, ingresos declarados).')
bullet('Ingresos Públicos › Val. Vehic.: catálogo de valuación vehicular (código de modelo, año, valor).')
para('En cada una, probá el ciclo completo: crear, editar y eliminar un registro.', italic=True, color=GREY)

# ════════════════════════════════════════════════════════════════════════
h1('8. Comunicación (CRUD de mensajes)')
step(1, 'Ir a Comunicación › solapa Mensajes › “+ Nuevo”.')
step(2, 'Completá Identificador, Título, Cuerpo y elegí Tipo / Canal / Prioridad / Estado (combos).')
step(3, 'Guardar → aparece en la tabla. Editá el título y guardá. Luego Eliminá.')
para('Resultado esperado: el alta, la edición y la baja funcionan sin error.', bold=True)

# ════════════════════════════════════════════════════════════════════════
h1('9. Verificación técnica por API (opcional)')
para('Para QA técnico, se puede validar la salud y los endpoints con curl:')
code('curl http://101.44.13.252/health\n'
     'curl http://101.44.13.252/api/seguridad/health\n'
     '# token:\n'
     'curl -X POST http://101.44.13.252/api/seguridad/auth/token \\\n'
     '  -H "Content-Type: application/json" \\\n'
     '  -d \'{"username":"admin","password":"Admin@2024!"}\'')

# ════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h1('Apéndice A — Sembrar la fórmula de demo (SQL)')
para('Lo ejecuta un administrador en el servidor, una sola vez, para habilitar el cálculo del ejemplo:')
code('cd ~/cheyenne\n'
     "docker compose exec -T postgres sh -c 'psql -U \"$POSTGRES_USER\" -d \"$POSTGRES_DB\"' <<'SQL'\n"
     "DELETE FROM emisiones_formula_tasa WHERE tipo_tributo='inmuebles' AND ttas_tasa=1;\n"
     "INSERT INTO emisiones_formula_tasa\n"
     " (tipo_tributo,ttas_tasa,ttas_subtasa,fort_numero,fort_orden,fort_descripcion,\n"
     "  fort_condicion,fort_a_cancelar_1,fort_a_pagar_1,activo,created_at)\n"
     "VALUES\n"
     " ('inmuebles',1,0,1,1,'ABL demo','#I_VALUACION(0) > 0',\n"
     "  '#REDONDEO(#I_VALUACION(0) * 0.012 + #I_SUPERFICIE(1,1) * 15, 2)',\n"
     "  '#REDONDEO((#I_VALUACION(0) * 0.012 + #I_SUPERFICIE(1,1) * 15) * 0.9, 2)',\n"
     "  true, now());\nSQL")
para('Las valuaciones y superficies del punto 3 se pueden cargar desde la interfaz '
     '(no requieren SQL).', italic=True, color=GREY)

h1('Apéndice B — Referencia rápida')
table(['Concepto', 'Detalle'], [
    ['URL', 'http://101.44.13.252/'],
    ['Login', 'admin / Admin@2024!'],
    ['Tributo de prueba', 'inmuebles · período 2026'],
    ['Fórmula ABL', 'valor × 1,2% + superficie × 15 ; a pagar con 10% de descuento'],
    ['Liquidaciones esperadas', '5400/4860 y 3150/2835 (total a pagar 7695)'],
    ['Deuda en cuenta corriente', '8550 (suma de “a cancelar”)'],
    ['Comprobantes', '2 recibos con código de barras'],
])

doc.add_paragraph()
foot = doc.add_paragraph(); foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = foot.add_run('Cheyenne · Guía de pruebas del flujo completo'); r.italic = True; r.font.color.rgb = GREY; r.font.size = Pt(9)

import os
out = os.path.join(os.path.dirname(__file__), 'Cheyenne-Guia-de-Pruebas.docx')
doc.save(out)
print('OK ->', out)
